from decouple import config
from langchain.chat_models import ChatOpenAI
from langchain.schema import HumanMessage, SystemMessage

from PyPDF2 import PdfReader
from io import BytesIO

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert
from datetime import datetime

from collections import defaultdict

chat_history = defaultdict(list)

# OpenAI API 키 설정
contracts_model = ChatOpenAI(model_name='gpt-3.5-turbo', temperature=0.5, openai_api_key=config('OPENAI_API_KEY'))
contracts_sys = SystemMessage(content="\
        저는 차용증을 작성해드리는 행정사입니다.\
        만약 사용자가 pdf양식을 업로드한다면 response에 양식을 보내주지 말고, 참고만 해줘\
        사용자가 입력한 내용을 차용증 형식에 맞게 정리하고, 필요한 정보가 누락되었을 경우 사용자에게 추가 정보를 요청할 것입니다.\
        만약 사용자가 원한다면 명시하지 않을 수 있지만 권장하지 않습니다.\
        차용증에는 다음과 같은 필수 정보가 포함되어야 합니다:\
        1. 차용금액 (숫자와 한글 병기)\
        2. 빌린 날짜\
        3. 이자 - 이자를 설정하려면 이자율 명시(법정 최고 이자 연 20%를 초과할 수 없습니다.)\
        4. 상환 방식 - 원금만기일시상환, 원금 균등상환, 원리금 균등상환, 원리금 일부 상환 후 만기일에 잔액 상환 등 방법이 정해지면 \
        5.   \
        4. 채무자 개인정보: 이름, 주소, 주민등록번호, 전화번호\
        5. 채권자 개인정보: 이름, 주소, 주민등록번호, 전화번호\
        6. 추가입력(없으면 안 적어도 됩니다.)\
                    특약사항: 연체 시 조항 등\
                    연대보증인: 연대보증인 개인정보와 책임\
        이자와 상환예정일자는 차용의 성격에 따라 명시하지 않을 수 있지만 권장합니다. 또한, 상환 종류에 대한 설명도 필요할 수 있습니다 (예: 원금만기일시상환, 원금 균등상환, 원리금 균등상환 등)\
        계약서 작성을 위해서 정확한 정보를 입력해주세요.\
        차용증 작성에 대한 내용만 입력해주세요. 다른 주제의 내용을 입력하면 이에 대한 답변을 드리지 않습니다.\
    ")


# 계약 작성 시작할 때 사용하는 서비스 메서드
# 대화 내용 저장
def get_initial_contract_response(user_id: int, service_id: int, chat: str) -> str:
    contracts_msg = HumanMessage(content=chat)

    response = contracts_model([contracts_sys, contracts_msg]).content
    conversation = {'user': chat, 'bot': response}
    
    save_history(user_id, service_id, conversation)

    return response


# 계약 작성 중 대화할 때 LLM 답변 요청하는 서비스 메서드
# 대화 내용 저장
def get_ongoing_contract_response(user_id: int, service_id: int, chat: str, history: dict) -> str:
    # chat_history에서 이전 대화 내용 복원
    history_key = make_history_key(user_id, service_id)
    previous_conversations = chat_history[history_key]
    conversations = ' '.join([item['user'] + item['bot'] for item in previous_conversations]) + chat
    contracts_msg = HumanMessage(content=conversations)

    response = contracts_model([contracts_sys, contracts_msg]).content
    conversation = {'user': chat, 'bot': response}

    save_history(user_id, service_id, conversation)

    return response


# 계약 작성 시작할 때 사용하는 서비스 메서드(PDF 업로드 기능)
# 대화 내용 저장(get_contract_response에 포함)
def get_initial_contract_response_with_pdf(user_id: int, service_id: int, file_contents: bytes, chat: str) -> str:
    pdf_text = convert_file_to_text(file_contents)
    request_text = pdf_text + chat

    return get_initial_contract_response(user_id, service_id, request_text)


# 계약 작성 중 사용하는 서비스 메서드
def get_additional_request_response(user_id: int, service_id: int, chat: str) -> str:
    history_key = make_history_key(user_id, service_id)

    # chat_history에 history_key로 저장된 내용이 있다면! (추후 히스토리 제한)
    conversations = chat_history[history_key]
    response = get_ongoing_contract_response(user_id, service_id, chat, conversations)

    return response

def convert_to_pdf(contract_info: ContractInfo) -> bytes:
    word = contractInfo_to_word(contract_info)
    html = word.to_html()

    # HTML을 PDF로 변환
    pdf = pdfkit.from_string(html, False)

    # BytesIO 객체로 변환하여 반환
    pdf_stream = BytesIO(pdf)
    return pdf_stream

# ContractInfo to word
# contract_info가 입력됐다고 가정하고 작성
# 프로토타입 모델
def contractInfo_to_word(contract_info: ContractInfo) -> Document:
    doc = Document()

    # 문서 제목
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('차용증')
    run.bold = True
    run.font.size = Pt(24)

    # 차용금액
    doc.add_paragraph(f'원({contract_info.loanDetails.loanAmountInKorean})')

    # 빌린 날짜 및 이자, 원금 상환일
    repayment_condition = contract_info.repaymentDetails.repaymentCondition or '-'
    repayment_method = contract_info.repaymentDetails.repaymentMethod or '-'
    maturity_date = contract_info.repaymentDetails.maturityDate or '-'
    
    doc.add_paragraph(f'위 금액을 채무자가 채권자로부터 {contract_info.loanDetails.borrowedDate} 틀림없이 빌렸습니다. '
                      f'이자는 연 {repayment_condition}으로 하여 매월 {repayment_method}일까지 갚겠으며, '
                      f'원금은 {maturity_date}까지 채권자에게 갚겠습니다.')

    # 채무자 정보
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    p.add_run(f'채무자 이름 : {contract_info.debtorInfo.name} (서명 또는 인)')
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    p.add_run(f'(빌리는 사람) 주소 : {contract_info.debtorInfo.address}')
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    p.add_run(f'주민등록번호 : {contract_info.debtorInfo.ssn}')
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    p.add_run(f'전화번호 : {contract_info.debtorInfo.contact}')

    # 채권자 정보
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    p.add_run(f'채권자 이름 : {contract_info.creditorInfo.name} 귀하')
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    p.add_run(f'(빌려주는 사람) 주소 : {contract_info.creditorInfo.address}')

    # 작성 날짜
    today = datetime.today().strftime('%Y년 %m월 %d일')
    doc.add_paragraph(today)

    # 파일 저장
    doc.save('차용증.docx')

    return doc

# Util

def convert_file_to_text(file_contents: bytes) -> str:
    pdf_file = BytesIO(file_contents)
    reader = PdfReader(pdf_file)
    
    text = ""
    for page_num in range(len(reader.pages)):
        text += reader.pages[page_num].extract_text()
        
    return text

def save_history(user_id: int, service_id: int, conversation: dict):
    history_key = make_history_key(user_id, service_id)
    chat_history[history_key].append(conversation)

def make_history_key(user_id: int, service_id: int) -> str:

    return str(user_id) + ':' + str(service_id)